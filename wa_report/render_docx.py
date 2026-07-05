"""Render a Report to a Word .docx file using python-docx.

Layout: one section per calendar day (header = the date), with a short clinical
summary (comfort / humidifier / water-trap) followed by the whole chat for that
day rendered chronologically, images inline, each item stamped with time only.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from . import media
from .extract import BUCKET_TITLES, COMFORT, HUMIDIFIER, WATER_TRAP
from .grouping import hm, transcript
from .report import Report

# Set to False to put ONLY the chat in the Word file (no Patient comfort /
# Humidifier / Water-trap summary block under each date).
SHOW_DAILY_SUMMARY = False

_BUCKET_ORDER = [COMFORT, HUMIDIFIER, WATER_TRAP]
_ACCENT = RGBColor(0x0A, 0x6E, 0xBD)
_MUTED = RGBColor(0x5B, 0x6B, 0x7B)
_PHOTOS_PER_ROW = 4
_PHOTO_WIDTH = Inches(1.45)
_CENTER = WD_ALIGN_PARAGRAPH.CENTER


def _label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = _MUTED
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def _keep_row_together(row):
    """Stop a table row from splitting across a page so an image and its time
    caption never end up on different pages (image on page 1, stamp on page 2)."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _flush_images(doc, imgs, cache):
    """Render a run of consecutive images as a centered grid with time captions."""
    if not imgs:
        return
    n = len(imgs)
    nrows = (n + _PHOTOS_PER_ROW - 1) // _PHOTOS_PER_ROW
    grid = doc.add_table(rows=nrows, cols=_PHOTOS_PER_ROW)
    for row in grid.rows:
        _keep_row_together(row)
    for i, (path, dt) in enumerate(imgs):
        cell = grid.cell(i // _PHOTOS_PER_ROW, i % _PHOTOS_PER_ROW)
        para = cell.paragraphs[0]
        para.alignment = _CENTER
        tmp = cache.temp_file(path)
        if tmp is not None:
            try:
                para.add_run().add_picture(str(tmp), width=_PHOTO_WIDTH)
            except Exception:
                para.add_run(path.name)
        cpara = cell.add_paragraph()
        cpara.alignment = _CENTER
        crun = cpara.add_run(hm(dt))
        crun.font.size = Pt(8)
        crun.font.color.rgb = _MUTED


def _chat_line(doc, t, speaker, text, *, muted=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    r = para.add_run(f"{t}  ")
    r.font.size = Pt(8)
    r.font.color.rgb = _MUTED
    w = para.add_run(f"{speaker}: ")
    w.bold = True
    w.font.size = Pt(9.5)
    tr = para.add_run(text)
    if muted:
        tr.italic = True
        tr.font.color.rgb = _MUTED
        tr.font.size = Pt(9)


def render_docx(report: Report, out_path: Path, cache: media.MediaCache) -> Path:
    doc = Document()
    doc.styles["Normal"].font.name = "Segoe UI"
    doc.styles["Normal"].font.size = Pt(10.5)

    # Header
    doc.add_heading(report.hospital, level=0)
    sub = doc.add_paragraph("Mechanical Ventilation Monitoring Report")
    sub.runs[0].font.color.rgb = _MUTED
    sub.runs[0].italic = True

    # Profile table
    rows = [
        ("Device", report.patient),
        ("Date In", report.date_in_str),
        ("Date Out", report.date_out_str),
        ("Duration", f"{report.duration_days} day(s)"),
        ("Group members", ", ".join(report.members) or "-"),
    ]
    ptab = doc.add_table(rows=0, cols=2)
    ptab.style = "Light List Accent 1"
    for k, v in rows:
        cells = ptab.add_row().cells
        cells[0].text = k
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = str(v)
    doc.add_paragraph()

    # One section per day
    for c in report.cycles:
        h = doc.add_heading(c.title, level=1)
        for run in h.runs:
            run.font.color.rgb = _ACCENT

        # Clinical summary (states noted that day) -- skipped when disabled
        if SHOW_DAILY_SUMMARY:
            for bucket in _BUCKET_ORDER:
                lines = c.comments.get(bucket)
                if lines:
                    _label(doc, BUCKET_TITLES[bucket])
                    for ln in lines:
                        doc.add_paragraph(ln, style="List Bullet")

        # Full chat, chronological, images inline (time only)
        img_run = []
        missing = 0
        for it in transcript(c):
            if it.kind == "image":
                p = media.resolve(it.name, it.source_dir)
                if p is not None:
                    img_run.append((p, it.dt))
                else:
                    missing += 1
                continue
            _flush_images(doc, img_run, cache)
            img_run = []
            _chat_line(doc, hm(it.dt), it.speaker, it.text,
                       muted=(it.kind == "media"))
        _flush_images(doc, img_run, cache)

        if missing:
            mp = doc.add_paragraph(
                f"{missing} image(s) referenced but not found in the export folder.")
            mp.runs[0].italic = True
            mp.runs[0].font.color.rgb = _MUTED
            mp.runs[0].font.size = Pt(9)

    doc.save(str(out_path))
    return out_path
